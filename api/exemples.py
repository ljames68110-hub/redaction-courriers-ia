from fastapi import APIRouter, HTTPException, UploadFile, File, Form
from typing import List
from core.db import SessionLocal
from models.orm_models import Courrier, Detenu
from ia import rag
import uuid
import io

router = APIRouter()

def _extract_text_from_docx(content: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Impossible de lire le fichier Word : {e}")

@router.post("/importer")
async def importer_courrier_exemple(
    user_id: str = Form(...),
    detenu_id: str = Form(...),
    type_courrier: str = Form(...),
    fichier: UploadFile = File(...)
):
    db = SessionLocal()
    det = db.query(Detenu).filter(Detenu.id == detenu_id, Detenu.user_id == user_id).first()
    if not det:
        raise HTTPException(status_code=404, detail="Détenu introuvable")
    if not fichier.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Seuls les fichiers .docx sont acceptés")

    content = await fichier.read()
    texte = _extract_text_from_docx(content)
    if not texte.strip():
        raise HTTPException(status_code=400, detail="Le fichier est vide ou illisible")

    from api.courriers import _get_fernet
    f = _get_fernet()
    token = f.encrypt(texte.encode("utf-8"))

    courrier_id = str(uuid.uuid4())
    courrier = Courrier(
        id=courrier_id,
        user_id=user_id,
        detenu_id=detenu_id,
        type=type_courrier,
        motif=f"[EXEMPLE IMPORTÉ] {fichier.filename}",
        contenu_chiffre=token.decode("utf-8"),
        tags="exemple"
    )
    db.add(courrier)
    db.commit()

    try:
        rag.index_courrier(courrier_id=courrier_id, detenu_id=detenu_id, text=texte)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erreur indexation RAG : {e}")

    return {
        "status": "ok",
        "courrier_id": courrier_id,
        "nom_fichier": fichier.filename,
        "nb_caracteres": len(texte)
    }


@router.post("/importer-masse")
async def importer_masse(
    user_id: str = Form(...),
    fichiers: List[UploadFile] = File(...)
):
    """
    Import en masse de fichiers .docx sans détenu obligatoire.
    Les courriers sont indexés comme exemples globaux (detenu_id='global').
    """
    from api.courriers import _get_fernet

    resultats = []
    erreurs = []

    for fichier in fichiers:
        if not fichier.filename.lower().endswith(".docx"):
            erreurs.append({"fichier": fichier.filename, "erreur": "Format non supporté (pas .docx)"})
            continue
        try:
            content = await fichier.read()
            # Extraire texte
            from docx import Document
            doc = Document(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            texte = "\n".join(paragraphs)

            if not texte.strip():
                erreurs.append({"fichier": fichier.filename, "erreur": "Fichier vide ou illisible"})
                continue

            # Détecter le type depuis le nom du fichier
            type_courrier = _detecter_type(fichier.filename)

            # Sauvegarder en base
            db = SessionLocal()
            f = _get_fernet()
            token = f.encrypt(texte.encode("utf-8"))

            courrier_id = str(uuid.uuid4())
            courrier = Courrier(
                id=courrier_id,
                user_id=user_id,
                detenu_id=None,
                type=type_courrier,
                motif=f"[BASE] {fichier.filename}",
                contenu_chiffre=token.decode("utf-8"),
                tags="exemple"
            )
            db.add(courrier)
            db.commit()

            # Indexer dans FAISS
            rag.index_courrier(
                courrier_id=courrier_id,
                detenu_id="global",
                text=texte
            )

            resultats.append({
                "fichier": fichier.filename,
                "type_detecte": type_courrier,
                "nb_caracteres": len(texte),
                "status": "ok"
            })

        except Exception as e:
            erreurs.append({"fichier": fichier.filename, "erreur": str(e)})

    return {
        "importes": len(resultats),
        "erreurs": len(erreurs),
        "detail_ok": resultats,
        "detail_erreurs": erreurs
    }


def _detecter_type(nom_fichier: str) -> str:
    """Détecte automatiquement le type de courrier depuis le nom du fichier."""
    nom = nom_fichier.lower()
    if any(x in nom for x in ["parloir", "visite"]):
        return "demande de parloir"
    elif any(x in nom for x in ["liberté", "liberation", "liberation"]):
        return "demande de mise en liberté"
    elif any(x in nom for x in ["jap", "application des peines"]):
        return "requête JAP"
    elif any(x in nom for x in ["transfert"]):
        return "demande de transfert"
    elif any(x in nom for x in ["médical", "medical", "sante", "santé", "ucsa"]):
        return "signalement médical"
    elif any(x in nom for x in ["greffe"]):
        return "demande au greffe"
    elif any(x in nom for x in ["direction"]):
        return "demande à la direction"
    elif any(x in nom for x in ["recours", "gracieux"]):
        return "recours gracieux"
    elif any(x in nom for x in ["discipline"]):
        return "recours disciplinaire"
    elif any(x in nom for x in ["spip"]):
        return "courrier SPIP"
    elif any(x in nom for x in ["uvf", "familiale"]):
        return "demande UVF"
    elif any(x in nom for x in ["amenagement", "aménagement"]):
        return "demande d'aménagement de peine"
    else:
        return "courrier administratif"


@router.get("/exemples")
def liste_exemples(user_id: str):
    db = SessionLocal()
    courriers = db.query(Courrier).filter(
        Courrier.user_id == user_id,
        Courrier.tags == "exemple"
    ).order_by(Courrier.created_at.desc()).all()

    return [
        {
            "id": c.id,
            "detenu_id": c.detenu_id,
            "type": c.type,
            "motif": c.motif,
            "created_at": c.created_at.isoformat() if c.created_at else None
        }
        for c in courriers
    ]

