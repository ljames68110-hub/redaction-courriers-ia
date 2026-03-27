from fastapi import APIRouter, HTTPException
from models.pydantic_schemas import CourrierGenerate, CourrierOut
from core.db import SessionLocal
from models.orm_models import Courrier, Detenu
from ia.model_loader import generate_courrier_text
from ia import rag
import uuid
from core.config import settings
from cryptography.fernet import Fernet
import base64

router = APIRouter()

def _get_fernet():
    key = base64.urlsafe_b64encode(settings.SQLCIPHER_KEY.encode("utf-8").ljust(32, b"\0"))
    return Fernet(key)

@router.get("/liste")
def liste_courriers(user_id: str):
    """Liste tous les courriers générés pour un utilisateur (hors exemples)."""
    db = SessionLocal()
    courriers = db.query(Courrier).filter(
        Courrier.user_id == user_id,
        Courrier.tags != "exemple"
    ).order_by(Courrier.created_at.desc()).all()

    f = _get_fernet()
    result = []
    for c in courriers:
        # Déchiffrer le contenu
        try:
            contenu = f.decrypt(c.contenu_chiffre.encode("utf-8")).decode("utf-8")
        except Exception:
            contenu = c.contenu_chiffre  # fallback texte brut

        # Récupérer infos détenu si lié
        det_nom = ""
        if c.detenu_id:
            det = db.query(Detenu).filter(Detenu.id == c.detenu_id).first()
            if det:
                det_nom = f"{det.nom} {det.prenom}"

        result.append({
            "id": c.id,
            "type": c.type,
            "motif": c.motif,
            "contenu": contenu,
            "detenu_id": c.detenu_id,
            "detenu_nom": det_nom,
            "created_at": c.created_at.isoformat() if c.created_at else None
        })
    return result

@router.post("/export-word")
def export_word(data: dict):
    """Génère un fichier .docx depuis un texte et le retourne en téléchargement."""
    from fastapi.responses import StreamingResponse
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    texte = data.get("contenu", "")
    nom_fichier = data.get("nom_fichier", "Courrier.docx")

    doc = Document()

    # Marges standard courrier
    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2)

    # Ajouter chaque ligne comme paragraphe
    for ligne in texte.split('\n'):
        p = doc.add_paragraph()
        run = p.add_run(ligne)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        # Gras pour la ligne Objet
        if ligne.strip().lower().startswith('objet'):
            run.bold = True
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        # Ligne vide = espace
        if not ligne.strip():
            p.paragraph_format.space_after = Pt(6)

    # Sauvegarder en mémoire
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{nom_fichier}"'}
    )

@router.post("/generer", response_model=CourrierOut)
def generer_courrier(payload: CourrierGenerate):
    db = SessionLocal()
    detenu_data = None
    detenu_id = payload.detenu_id

    if payload.detenu_id:
        det = db.query(Detenu).filter(Detenu.id == payload.detenu_id).first()
        if not det:
            raise HTTPException(status_code=404, detail="Detenu introuvable")
        detenu_data = {
            "id": det.id,
            "nom": det.nom,
            "prenom": det.prenom,
            "numero_ecrou": det.numero_ecrou,
            "cellule": det.cellule,
            "batiment": det.batiment,
            "etablissement": det.etablissement,
            "date_naissance": det.date_naissance
        }

    elif payload.detenu_temp:
        detenu_data = payload.detenu_temp.dict()
        if payload.save_detenu:
            new_id = str(uuid.uuid4())
            det = Detenu(
                id=new_id,
                user_id=payload.user_id,
                nom=detenu_data["nom"],
                prenom=detenu_data["prenom"],
                numero_ecrou=detenu_data["numero_ecrou"],
                cellule=detenu_data["cellule"],
                batiment=detenu_data["batiment"],
                etablissement=detenu_data.get("etablissement"),
                date_naissance=detenu_data.get("date_naissance")
            )
            db.add(det)
            db.commit()
            detenu_id = new_id
            detenu_data["id"] = new_id

    generated_text, metadata = generate_courrier_text(
        detenu=detenu_data,
        detenu_id=detenu_id or (detenu_data.get("id") if detenu_data else None),
        type_courrier=payload.type_courrier,
        motif=payload.motif,
        ton=payload.ton
    )

    f = _get_fernet()
    token = f.encrypt(generated_text.encode("utf-8"))

    courrier_id = str(uuid.uuid4())
    courrier = Courrier(
        id=courrier_id,
        user_id=payload.user_id,
        detenu_id=detenu_id,
        type=payload.type_courrier,
        motif=payload.motif,
        contenu_chiffre=token.decode("utf-8"),
        tags=None
    )
    db.add(courrier)
    db.commit()

    try:
        rag.index_courrier(
            courrier_id=courrier_id,
            detenu_id=detenu_id or (detenu_data.get("id") if detenu_data else ""),
            text=generated_text
        )
    except Exception as e:
        print("Erreur indexation RAG:", e)

    return CourrierOut(courrier_id=courrier_id, contenu=generated_text, metadata=metadata)
